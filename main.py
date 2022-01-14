from PIL import Image
from pathlib import Path
import zlib
import numpy
from pikepdf import Pdf, PdfImage, Name
import cv2
import zipfile
from docx import Document
from docx.shared import Inches, Cm, Pt
from docxtpl import DocxTemplate
import io
import fire


def remove_watermark_from_image(img: numpy.ndarray) -> numpy.ndarray:
    """
    Remove watermark from open cv image
    :param img:
    :return: image without watermark
    """
    # convert image to hsv colorspace
    hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
    h, s, v = cv2.split(hsv)
    # threshold saturation image
    thresh1 = cv2.threshold(s, 92, 255, cv2.THRESH_BINARY)[1]
    # threshold value image and invert
    thresh2 = cv2.threshold(v, 128, 255, cv2.THRESH_BINARY)[1]
    thresh2 = 255 - thresh2
    # combine the two threshold images as a mask
    mask = cv2.add(thresh1, thresh2)
    # use mask to remove lines in background of input
    result = img.copy()
    result[mask == 0] = (255, 255, 255)
    return result


def generate_output_path(input_path: Path) -> Path:
    """
    Generate output_path from input_path
    :param input_path:
    :return:
    """
    return input_path.parent / (input_path.stem + f"_generated{input_path.suffix}")


def remove_watermark_from_pdf(input_file: Path, output_file: Path) -> str:
    """
    Remove watermark from pdf and save to output_file
    :param input_file:
    :param output_file:
    :return:
    """
    pdf = Pdf.open(input_file)
    for page in pdf.pages:
        for image_key in page.images.keys():
            raw_image = page.images[image_key]
            pdf_image = PdfImage(raw_image)
            pil_image = pdf_image.as_pil_image()
            # noinspection PyTypeChecker
            opencv_image = cv2.cvtColor(numpy.array(pil_image), cv2.COLOR_RGB2BGR)
            opencv_image = remove_watermark_from_image(opencv_image)
            pil_image = Image.fromarray(cv2.cvtColor(opencv_image, cv2.COLOR_BGR2RGB))
            raw_image.write(zlib.compress(pil_image.tobytes()), filter=Name("/FlateDecode"))
    pdf.save(output_file)
    return str(output_file)


def remove_watermark_from_docx(input_file: Path, output_file: Path) -> str:
    """
    Remove watermark from docx and save to output_file
    :param input_file:
    :param output_file:
    :return:
    """
    # TODO: use tmplx instead of docx
    z = zipfile.ZipFile(input_file)
    # print list of valid attributes for ZipFile object
    # print(dir(z))
    # print all files in zip archive
    all_files = z.namelist()
    # print(all_files)
    # get all files in word/media/ directory
    images = list(filter(lambda x: x.startswith('word/media/'), all_files))
    print(images)
    # print(input_file.absolute())
    doc_tmpl = DocxTemplate(input_file.absolute())
    coc_tmpl = doc_tmpl.render({})
    # print(doc_tmpl.get_docx())
    # print(doc_tmpl.get_xml())
    # print(doc_tmpl.get_pic_map())
    for image in images:
        # placeholder_image = Path(image).stem + Path(image).suffix
        # print('placeholder_image', placeholder_image)
        with z.open(image) as f:
            pil_image = Image.open(f)
            # noinspection PyTypeChecker
            opencv_image = cv2.cvtColor(numpy.array(pil_image), cv2.COLOR_RGB2BGR)
            opencv_image = remove_watermark_from_image(opencv_image)
            pil_image = Image.fromarray(cv2.cvtColor(opencv_image, cv2.COLOR_BGR2RGB))
            image_file_tmp = io.BytesIO()
            pil_image.save(image_file_tmp, format="PNG")
            print('image_file_tmp', image_file_tmp)
            # doc_tmpl.replace_media(image, image_file_tmp)
            doc_tmpl.replace_pic(image, image_file_tmp)
    doc_tmpl.save(output_file)
    z.close()
    return str(output_file)
    # tpl.replace_pic('dummy_header_pic.jpg','header_pic_i_want.jpg')
    # removing margins
    document = Document()
    margin = Cm(0.0)
    for section in document.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
    # document.add_page_break()
    nb_images = len(images)
    # open an image and save it
    for i, image in enumerate(images):
        with z.open(image) as f:
            pil_image = Image.open(f)
            # noinspection PyTypeChecker
            opencv_image = cv2.cvtColor(numpy.array(pil_image), cv2.COLOR_RGB2BGR)
            opencv_image = remove_watermark_from_image(opencv_image)
            pil_image = Image.fromarray(cv2.cvtColor(opencv_image, cv2.COLOR_BGR2RGB))
            # add image to document with full page width and height
            # document.add_picture(img, width=Cm(21.0), height=Cm(29.7))
            image_file_tmp = io.BytesIO()
            # pil_image.save(image_file_tmp, format=pil_image.format) # .format lost because of conversion with cv2
            pil_image.save(image_file_tmp, format="PNG")
            # document.add_picture(image_file_tmp, width=Pt(pil_image.size[0]), height=Pt(pil_image.size[1]))
            # document.add_picture(image_file_tmp, width=Cm(21.0), height=Cm(29.7))
            # document.add_picture(image_file_tmp, height=Cm(29.7))
            # document.add_picture(image_file_tmp, width=Cm(21.0))
            document.add_picture(image_file_tmp, width=Inches(8.5), height=Inches(11.0))
            if i < nb_images - 1:
                document.add_page_break()
    z.close()
    document.save(output_file)
    return str(output_file)


def main(input_file: str, output_file: str = None) -> str:
    """
    Entry point
    :param input_file:
    :param output_file:
    :return:
    """
    input_path = Path(input_file)
    if not input_path or not input_path.exists():
        raise FileNotFoundError(f"File {input_path} does not exist")
    # generate output path
    if not output_file:
        output_path = generate_output_path(input_path)
    else:
        output_path = Path(output_file)
    # remove output file if it exists
    if output_path.exists():
        output_path.unlink()
    # remove watermark
    if str(input_path.suffix).lower() == ".pdf":
        return remove_watermark_from_pdf(input_path, output_path)
    elif str(input_path.suffix).lower() == ".docx":
        return remove_watermark_from_docx(input_path, output_path)
    else:
        raise Exception(f"Unsupported file type: {input_path.suffix}")


if __name__ == "__main__":
    fire.Fire(main)
